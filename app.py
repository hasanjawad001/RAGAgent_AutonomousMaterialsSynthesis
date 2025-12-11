from pathlib import Path
import re
import fitz
import tiktoken
from openai import OpenAI
import numpy as np
import faiss
import streamlit as st
import pickle
import os
import openai
from langchain.text_splitter import RecursiveCharacterTextSplitter
import pycountry
import time
from langchain_community.vectorstores import FAISS 
from langchain_community.docstore.in_memory import InMemoryDocstore
from langchain_core.documents import Document
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
import pandas as pd
import io
import base64
import mimetypes
from docx import Document as DocxDocument
from langchain_community.retrievers import BM25Retriever
from langchain.retrievers import EnsembleRetriever, MultiQueryRetriever
from langchain.retrievers.contextual_compression import ContextualCompressionRetriever
from langchain.retrievers.document_compressors import LLMChainExtractor
from langchain.retrievers.document_compressors import CrossEncoderReranker
from langchain_community.cross_encoders import HuggingFaceCrossEncoder
from langchain_community.document_transformers import LongContextReorder
from langchain_experimental.graph_transformers import LLMGraphTransformer
import pickle
import networkx as nx
from langchain_core.retrievers import BaseRetriever
# from langchain.globals import set_llm_cache
# from langchain_community.cache import SQLiteCache
import json
import subprocess
import yaml
from lightrag import LightRAG, QueryParam
from lightrag.llm.openai import gpt_4o_mini_complete, gpt_4o_complete, openai_embed
from lightrag.utils import EmbeddingFunc
import asyncio
from lightrag.kg.shared_storage import initialize_pipeline_status
import uuid
import time
import nest_asyncio

################################
# Globals / Config
################################

# set_llm_cache(SQLiteCache(database_path=".cache.db"))
st.set_page_config(page_title="RAG Agent", layout="centered")

enc = tiktoken.get_encoding("cl100k_base")

d_em2dim = {
    "text-embedding-3-small": 1536,
    "text-embedding-3-large": 3072
}
TOKENS_PER_CHUNK = 300
WORDS_PER_CHUNK_OVERLAP = int(TOKENS_PER_CHUNK / 5)  # ~20%
top_k_textKBfaiss = 50
top_k_textKBbm = 50
top_k_addKBfaiss = 25
top_k_textKBgraph = 25
stream_delay = 0.08


################################
# Helpers 
################################

def dequote_path(p: str) -> str:
    if not p:
        return p
    p = p.strip()
    if len(p) >= 2 and p[0] == p[-1] and p[0] in "\"'":
        p = p[1:-1]
    return p

def extract_text_from_pdf(path):
    doc = fitz.open(path)
    return "\n".join(page.get_text() for page in doc)

def remove_junk_sections(text, section_markers=None):
    if section_markers is None:
        section_markers = [
            "references",
            "acknowledgment", "acknowledgement", "acknowledgments", "acknowledgements",
            "author information", "author contribution", "author contributions",
            "associated content",
        ]
    pattern = re.compile(rf"^\s*[\.\-\u25A0\u2022]*\s*({'|'.join(section_markers)})", re.IGNORECASE | re.MULTILINE)
    match = pattern.search(text)
    return text[:match.start()] if match else text

def remove_junk_lines(text, junk_patterns=None):
    if junk_patterns is None:
        junk_patterns = [
            "doi:", "et al.", "https://", "http://", ".org", ".com", "conflict of interest", "bio:", "funding",
            "journal", "citation", "cc-by", "preprint", "arxiv",
            "license", "open access",
            "submitted to", "peer review", "double-blind", "published",
            "copyright", "all rights reserved",
            ## "figure", "table",
            "correspondence should be addressed to",
            "authors contributed equally",
            "this manuscript has been authored by",
            "contract no",
            "supporting information",
            "read online",
            "received:",
            "revised:",
            "accepted:",
            "cite this:",
            "©",
            "download",
            "public access plan",
            "department of",
            "university",
            "national laboratory",
            "laboratory",
            "government",
        ]
    countries = [country.name.lower() for country in pycountry.countries]
    lines = text.splitlines()
    cleaned = []
    for line in lines:
        lower = line.lower()
        if any(pat in lower for pat in junk_patterns):
            continue
        if "*" in line or "†" in line or "∇" in line:
            continue
        if any(re.search(rf"\b{re.escape(c)}\b", lower) for c in countries):
            continue
        cleaned.append(line)
    return "\n".join(cleaned)

def chunk_text(text, max_tokens=None, tokenizer=None):
    if max_tokens is None:
        max_tokens = TOKENS_PER_CHUNK
    if tokenizer is None:
        tokenizer = enc
    words = text.split()
    chunks = []
    current = []
    for word in words:
        current.append(word)
        test_chunk = " ".join(current)
        if len(tokenizer.encode(test_chunk)) > max_tokens:
            current.pop()
            chunks.append(" ".join(current))
            current = [word]
    if current:
        chunks.append(" ".join(current))
    return chunks

def chunk_text2(text, max_tokens=None, tokenizer=None, overlap=None):
    # Sliding-window with word overlap to keep coherence
    if max_tokens is None:
        max_tokens = TOKENS_PER_CHUNK
    if tokenizer is None:
        tokenizer = enc
    if overlap is None:
        overlap = WORDS_PER_CHUNK_OVERLAP
    words = text.split()
    chunks = []
    current = []
    i = 0
    while i < len(words):
        word = words[i]
        current.append(word)
        test_chunk = " ".join(current)
        if len(tokenizer.encode(test_chunk)) > max_tokens:
            current.pop()
            chunks.append(" ".join(current))
            overlap_start = max(len(current) - overlap, 0)
            current = current[overlap_start:] + [word]
        i += 1
    if current:
        chunks.append(" ".join(current))
    return chunks

def estimate_embedding_cost(token_count, model="text-embedding-3-small"):
    price_per_million = {
        "text-embedding-3-small": 0.02,
        "text-embedding-3-large": 0.13,
        "text-embedding-ada-002": 0.10,
    }.get(model, 0.02)  # price per 1,000,000 tokens
    return (token_count / 1_000_000) * price_per_million

def _to_data_url(file_bytes, mime_type="image/png"):
    # Convert bytes -> data URL for multimodal image input
    b64 = base64.b64encode(file_bytes).decode("utf-8")
    return f"data:{mime_type};base64,{b64}"

def _parse_table_file(file_bytes, filename, max_rows=50, max_chars=20000):
    # Compact text extraction for tabular files
    try:
        if filename.lower().endswith(".csv"):
            df = pd.read_csv(io.BytesIO(file_bytes))
        else:  # .xlsx
            df = pd.read_excel(io.BytesIO(file_bytes), engine="openpyxl")
    except Exception as e:
        return f"[PARSE-ERROR {filename}: {e}]"
    buf = []
    buf.append(f"TABLE: {filename}")
    buf.append("COLUMNS: " + ", ".join(map(str, df.columns.tolist())))
    head = df.head(max_rows)
    buf.append("HEAD:")
    buf.append(head.to_csv(index=False))
    text = "\n".join(buf)
    if len(text) > max_chars:
        text = text[:max_chars] + "\n...[truncated]..."
    return text

def build_upload_bundle(uploaded_files, client, embedding_model, dimension):
    """
    Build in-memory FAISS for uploaded *text-like* files and collect images.
    Returns: (upload_db, text_meta, images)
      - upload_db: FAISS store for uploaded text chunks (or None if none)
      - text_meta: list of dicts for chunks
      - images: list of {"name": str, "data_url": str}
    """
    text_chunks = []
    text_meta = []
    images = []

    for uf in uploaded_files:
        name = uf.name
        mime = uf.type or mimetypes.guess_type(name)[0] or ""
        data = uf.getvalue()  # bytes

        # Images
        if name.lower().endswith((".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tif", ".tiff")):
            data_url = _to_data_url(data, mime or "image/png")
            images.append({"name": name, "data_url": data_url})
            continue

        # Text-like docs
        text = ""
        if name.lower().endswith(".pdf"):
            try:
                with fitz.open(stream=data, filetype="pdf") as doc:
                    text = "\n".join(page.get_text() for page in doc)
            except Exception as e:
                st.warning(f"⚠️ Could not read PDF {name}: {e}")
                continue
        elif name.lower().endswith(".txt"):
            text = data.decode("utf-8", errors="ignore")
        elif name.lower().endswith(".csv") or name.lower().endswith(".xlsx"):
            text = _parse_table_file(data, name)
        else:
            st.warning(f"⚠️ Unsupported file type: {name} (supported: pdf/txt/csv/xlsx + images)")
            continue

        # Clean + chunk
        text = remove_junk_sections(text)
        text = remove_junk_lines(text)
        chunks = chunk_text2(
            text, max_tokens=TOKENS_PER_CHUNK, tokenizer=enc, overlap=WORDS_PER_CHUNK_OVERLAP
        )
        for i, ch in enumerate(chunks):
            text_chunks.append(ch)
            text_meta.append({
                "source": f"uploaded/{name}",
                "chunk_id": i,
                "text": ch,
                "embedding_model": embedding_model
            })

    # Build FAISS for uploaded text
    upload_db = None
    if text_chunks:
        BATCH = 64
        embs = []
        for i in range(0, len(text_chunks), BATCH):
            batch = text_chunks[i:i + BATCH]
            resp = client.embeddings.create(input=batch, model=embedding_model)
            embs.extend([d.embedding for d in resp.data])

        emb_mat = np.array(embs, dtype="float32")
        index = faiss.IndexFlatL2(dimension)
        index.add(emb_mat)

        ids = [str(i) for i in range(len(text_meta))]
        docs_dict = {
            ids[i]: Document(
                page_content=text_meta[i]["text"],
                metadata={
                    "source": text_meta[i]["source"], 
                    "chunk_id": text_meta[i]["chunk_id"],
                    "original_content": text_meta[i]["text"], 
                }
            ) for i in range(len(text_meta))
        }
        docstore = InMemoryDocstore(docs_dict)
        index_to_docstore_id = {i: ids[i] for i in range(len(ids))}

        upload_db = FAISS(
            embedding_function=OpenAIEmbeddings(model=embedding_model,api_key=st.session_state.api_key),
            index=index,
            docstore=docstore,
            index_to_docstore_id=index_to_docstore_id,
        )

    return upload_db, text_meta, images

# @st.cache_resource
def get_lightrag_engine(working_dir, api_key, embedding_model, embedding_dim):
    """
    Initializes and caches the LightRAG engine.
    """
    if not os.path.exists(working_dir):
        os.makedirs(working_dir)
        
    rag = LightRAG(
        working_dir=working_dir,
        llm_model_func=lambda prompt, system_prompt=None, history_messages=[], **kwargs: gpt_4o_mini_complete(
            prompt, 
            system_prompt=system_prompt, 
            history_messages=history_messages, 
            api_key=api_key, 
            **kwargs
        ),        
        embedding_func=EmbeddingFunc(
            embedding_dim=embedding_dim, 
            max_token_size=8192,
            func=lambda texts: openai_embed(
                texts, 
                model=embedding_model, # <-- Chosen model name passed here
                api_key=api_key         # <-- API Key passed here
            ) 
        ),        
        llm_model_name="gpt-4o-mini",
    )    
    return rag

class StaticGraphRetriever(BaseRetriever):
    def _get_relevant_documents(self, query, *, run_manager=None):
        return graph_docs
    async def _aget_relevant_documents(self, query, *, run_manager=None):
        return graph_docs                

    
################################
# App UI
################################

st.markdown("## 📄 RAG Agent — Autonomous Synthesis")
st.divider()

# s1: API key
if "api_verified" not in st.session_state:
    st.session_state.api_verified = False
if "client" not in st.session_state:
    st.session_state.client = None
if not st.session_state.api_verified:
    with st.form("api_form"):
        api_key = st.text_input("🔑 Enter your OpenAI API key", type="password")
        verify_btn = st.form_submit_button("Verify")
        if verify_btn:
            try:
                client = OpenAI(api_key=api_key)
                _ = client.models.list()  # sanity check
                st.session_state.api_verified = True
                st.session_state.client = client
                st.session_state.api_key = api_key
                os.environ["OPENAI_API_KEY"] = api_key
                st.success("✅ API key verified!")
                st.rerun()
            except Exception as e:
                st.error(f"❌ Invalid API key: {e}")
    st.stop()
client = st.session_state.client

embeddings_obj = OpenAIEmbeddings(
    model=st.session_state.get("embedding_model", "text-embedding-3-small"),
    api_key=st.session_state.api_key,
)

# s2: KB setup
st.header("📂 Knowledge-Base Setup")
option = st.radio(
    "Choose how you want to set up the Knowledge-Base:",
    ("📚 Build a new Knowledge-Base", "📤 Load existing Knowledge-Base", "➕ Append existing Knowledge-Base"),
    index=0
)

if option == "📚 Build a new Knowledge-Base":
    st.session_state.embedding_model = st.selectbox(
        "🔍 Select your embedding model:",
        options=["text-embedding-3-small", "text-embedding-3-large"],
        index=0,
        help="This model is used for generating embeddings → impacts context matching"
    )
    st.session_state.dimension = d_em2dim[st.session_state.embedding_model]
    st.session_state.pdf_folder = dequote_path(st.text_input(
        "📁 Input Folder path for PDFs:",
        value='inputs',
        help="Path to get the PDFs as input to build the Knowledge-Base"
    ))
    st.session_state.knowledge_base_name = dequote_path(st.text_input(
        "🧠 Output FAISS index file path (.index)",
        value='outputs/test_index.index',
        help="Path to save the FAISS index"
    ))
    st.session_state.metadata_name = dequote_path(st.text_input(
        "📝 Output metadata file path (.pkl)",
        value='outputs/test_metadata.pkl',
        help="Path to save metadata for chunks"
    ))
    st.session_state.knowledge_graph_parent = dequote_path(st.text_input(
        "📁 Parent directory for Knowledge-Graph",
        value='outputs',
        help="Folder where the Knowledge-Graph pipeline will create files/folders and save related artifacts"
    ))
    if st.button("📚 Build Knowledge-Base"):
        if not os.path.isdir(st.session_state.pdf_folder):
            st.error("The provided folder path does not exist!")
            st.stop()
        pdf_files = list(Path(st.session_state.pdf_folder).glob("*.pdf"))
        if not pdf_files:
            st.error("No PDF files found in the selected folder!")
            st.stop()
        ##
        ## retrieving chunks, tokens
        total_chunks = 0
        total_tokens = 0
        total_cost = 0.0
        all_chunks = {}
        all_texts = {} 
        st.write("🔍 Processing PDFs...")
        progress_bar = st.progress(0)
        #
        for ipp, pdf_path in enumerate(pdf_files):
            try:
                text = extract_text_from_pdf(pdf_path)
                text = remove_junk_sections(text)
                text = remove_junk_lines(text)
                all_texts[pdf_path.name] = text
                chunks = chunk_text2(text, max_tokens=TOKENS_PER_CHUNK, tokenizer=enc, overlap=WORDS_PER_CHUNK_OVERLAP)
                token_count = sum(len(enc.encode(chunk)) for chunk in chunks)
                cost_emb = estimate_embedding_cost(token_count, model=st.session_state.embedding_model)
                cost = cost_emb
                total_chunks += len(chunks)
                total_tokens += token_count
                total_cost += cost
                all_chunks[pdf_path.name] = chunks
            except Exception as e:
                st.warning(f"⚠️ Failed on {pdf_path.name}: {e}")
            progress_percent = int((ipp + 1) / len(pdf_files) * 100)
            progress_bar.progress(progress_percent)
        #
        st.write(f"📦 Total chunks: {total_chunks:,}, Total tokens: {total_tokens:,}!")
        ## chunks, tokens -
        ## index, metadata
        st.write("📌 Embedding and indexing...")
        all_embeddings = []
        all_metadata = []
        chunk_count = 0
        progress_bar = st.progress(0)
        #
        for filename, chunks in all_chunks.items():
            for i, chunk in enumerate(chunks):
                try:
                    response = client.embeddings.create(input=chunk, model=st.session_state.embedding_model)
                    embedding = response.data[0].embedding
                    all_embeddings.append(embedding)
                    all_metadata.append({
                        "source": filename,
                        "chunk_id": i,
                        "text": chunk,
                        "embedding_model": st.session_state.embedding_model
                    })
                except Exception as e:
                    st.warning(f"⚠️ Embedding failed: {filename}, chunk_id: {i}, error: {e}")
                chunk_count += 1
                progress = int((chunk_count / total_chunks) * 100)
                progress_bar.progress(progress)
        #
        embedding_matrix = np.array(all_embeddings, dtype="float32")
        st.session_state.dimension = d_em2dim[st.session_state.embedding_model]
        index = faiss.IndexFlatL2(st.session_state.dimension)
        index.add(embedding_matrix)
        os.makedirs(os.path.dirname(st.session_state.knowledge_base_name), exist_ok=True)
        os.makedirs(os.path.dirname(st.session_state.metadata_name), exist_ok=True)
        faiss.write_index(index, st.session_state.knowledge_base_name)
        with open(st.session_state.metadata_name, "wb") as f:
            pickle.dump(all_metadata, f)
        st.success("✅ Knowledge-Base built and saved!")
        st.session_state.index = index
        st.session_state.metadata = all_metadata
        ## index, metadata -
        ## db
        metadata = all_metadata
        ids = [str(i) for i in range(len(metadata))]
        docs_dict = {
            ids[i]: Document(
                page_content=meta["text"],
                metadata={
                    "source": meta["source"], 
                    "chunk_id": meta["chunk_id"],
                    "original_content": meta["text"],                     
                }
            )
            for i, meta in enumerate(metadata)
        }
        docstore = InMemoryDocstore(docs_dict)
        index_to_docstore_id = {i: ids[i] for i in range(len(ids))}
        db = FAISS(
            embedding_function=embeddings_obj,
            index=index,
            docstore=docstore,
            index_to_docstore_id=index_to_docstore_id,
        )
        st.session_state.db = db
        ## db -
        ## bm
        all_documents = list(docstore._dict.values())  
        st.session_state.all_documents = all_documents
        st.session_state.bm25 = BM25Retriever.from_documents(all_documents, k=top_k_textKBbm)
        ## bm -
        ## knowledge-graph       
        if 'lightrag_engine' in st.session_state:
            kg_parent_dir = None
            kg_dir = None
            rag = None
            documents_to_ingest = None
            loop_build = None
            del st.session_state.lightrag_engine        
        try:
            kg_parent_dir = Path(st.session_state.get("knowledge_graph_parent", "outputs"))
            kg_dir = kg_parent_dir / "knowledge_graph"
            kg_dir.mkdir(parents=True, exist_ok=True)                
            rag = get_lightrag_engine(
                working_dir=str(kg_dir),
                api_key=st.session_state.api_key,
                embedding_model=st.session_state.get("embedding_model", "text-embedding-3-small"),
                embedding_dim=st.session_state.get("dimension", 1536),
            )
            documents_to_ingest = [f"{text}" for text in all_texts.values()]
            # documents_to_ingest = [f"DOC_ID:{uuid.uuid4().hex}\n{text}" for text in all_texts.values()]

            
            st.write("🕸️ Building Knowledge-Graph...")
            with st.spinner("🚀 Running indexing pipeline... (extraction, embedding, and graph creation)"):
                async def build_graph():
                    await rag.initialize_storages()
                    await initialize_pipeline_status()
                    await rag.ainsert(documents_to_ingest)
                    # await rag.finalize_storages()

                nest_asyncio.apply()                
                loop_build = asyncio.new_event_loop()
                asyncio.set_event_loop(loop_build)
                try:
                    loop_build.run_until_complete(build_graph())
                finally:
                    # loop_build.close()
                    pass

            st.session_state.lightrag_engine = rag            
            st.success("✅ Indexing completed! The Knowledge-Graph is ready!")
        except Exception as e:
            st.warning(f"⚠️ Knowledge-Graph build failed: {e}")
        ## knowledge-graph -

elif option == "📤 Load existing Knowledge-Base":
    index_path = st.text_input("🧠 Index file path (.index)", value='outputs/test_index.index')
    meta_path = st.text_input("📝 Metadata file path (.pkl)", value='outputs/test_metadata.pkl')
    knowledge_graph_parent = st.text_input(
        "🕸️ Parent directory for Knowledge-Graph",
        value='outputs',
        help="This folder must contain the 'knowledge_graph' subfolder where previously built Knowledge-Graph stored all graph artifacts."
    )

    if st.button("📤 Load Knowledge-Base"):
        if not os.path.isfile(index_path):
            st.error("Index file not found!")
            st.stop()
        if not os.path.isfile(meta_path):
            st.error("Metadata file not found!")
            st.stop()
        try:
            ## index, metadata, db
            index = faiss.read_index(index_path)
            with open(meta_path, "rb") as f:
                metadata = pickle.load(f)                

            ids = [str(i) for i in range(len(metadata))]
            docs_dict = {
                ids[i]: Document(
                    page_content=meta["text"],
                    metadata={
                        "source": meta["source"], 
                        "chunk_id": meta["chunk_id"],
                        "original_content": meta["text"],                    
                    }
                )
                for i, meta in enumerate(metadata)
            }
            docstore = InMemoryDocstore(docs_dict)
            index_to_docstore_id = {i: ids[i] for i in range(len(ids))}
            db = FAISS(
                embedding_function=embeddings_obj,
                index=index,
                docstore=docstore,
                index_to_docstore_id=index_to_docstore_id,
            )
            st.session_state.embedding_model = metadata[0].get("embedding_model", "text-embedding-3-small")
            st.session_state.dimension = d_em2dim[st.session_state.embedding_model]
            st.session_state.index = index
            st.session_state.metadata = metadata            
            st.success(f"✅ Knowledge-Base loaded successfully! Embedding model: `{st.session_state.embedding_model}`")                                    
            st.session_state.db = db
            ## index, metadata, db -
            # === loading for bm, graph ===
            all_documents = list(docstore._dict.values())
            st.session_state.all_documents = all_documents
            st.session_state.bm25 = BM25Retriever.from_documents(all_documents, k=top_k_textKBbm)            
            #            
            if 'lightrag_engine' in st.session_state:
                kg_parent_dir = None
                kg_dir = None
                rag = None
                loop_load = None
                del st.session_state.lightrag_engine                   
            kg_parent_dir = Path(knowledge_graph_parent)
            kg_dir = kg_parent_dir / "knowledge_graph"
            required_files = [
                "graph_chunk_entity_relation.graphml", 
                "kv_store_doc_status.json", "kv_store_full_entities.json", "kv_store_full_relations.json", 
                "kv_store_full_docs.json", "kv_store_text_chunks.json", "kv_store_entity_chunks.json", "kv_store_relation_chunks.json", 
                "vdb_entities.json", "vdb_relationships.json", "vdb_chunks.json",
            ]
            missing = [f for f in required_files if not (kg_dir / f).exists()]
            
            if missing:
                st.warning(f"⚠️ Knowledge-Graph is incomplete or corrupted.\nMissing: {missing}")
            else:
                rag = get_lightrag_engine(
                    working_dir=str(kg_dir),
                    api_key=st.session_state.api_key,
                    embedding_model=st.session_state.embedding_model,
                    embedding_dim=st.session_state.dimension,
                )                
                st.write("🕸️ Loading existing Knowledge-Graph...")
                with st.spinner("🚀 Running pipeline..."):
                    async def load_graph():
                        await rag.initialize_storages()                        
                        
                    nest_asyncio.apply()
                    loop_load = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop_load)
                    try:
                        loop_load.run_until_complete(load_graph())
                    finally:
                        # loop_load.close()  
                        pass
                st.session_state.lightrag_engine = rag            
                st.success("✅ Knowledge-Graph loaded successfully!")
            ##
            ## bm, grpah -
        except Exception as e:
            st.error(f"❌ Failed to load Knowledge-Base: {e}")
            st.stop()
            
elif option == "➕ Append existing Knowledge-Base":
    exist_index_path = st.text_input("🧠 Existing index file path (.index)", value="outputs/test_index.index")
    exist_meta_path  = st.text_input("📝 Existing metadata file path (.pkl)", value="outputs/test_metadata.pkl")
    exist_knowledge_graph_parent = st.text_input(
        "🕸️ Existing parent directory for Knowledge-Graph",
        value='outputs',
        help="This folder must contain the 'knowledge_graph' subfolder where previously built Knowledge-Graph stored all graph artifacts."
    )
    append_folder = dequote_path(st.text_input(
        "📁 Input Folder path for NEW PDFs to append:",
        value="inputs/new",
        help="Path to get the NEW PDFs to be appended"
    ))
    # out_index_path = dequote_path(st.text_input("🧠 Output (updated) FAISS index file path (.index)", value="outputs/test_index_appended.index"))
    # out_meta_path  = dequote_path(st.text_input("📝 Output (updated) metadata file path (.pkl)", value="outputs/test_metadata_appended.pkl"))
    out_index_path = exist_index_path
    out_meta_path  = exist_meta_path
    
    if st.button("➕ Append to Knowledge-Base"):
        # 0) Sanity checks
        if not os.path.isfile(exist_index_path):
            st.error("Existing FAISS index file not found!"); st.stop()
        if not os.path.isfile(exist_meta_path):
            st.error("Existing metadata file not found!"); st.stop()
        if not os.path.isdir(append_folder):
            st.error("Append folder does not exist!"); st.stop()
        if not os.path.isdir(exist_knowledge_graph_parent):
            st.error("Parent directory for Knowledge-Graph has not been found!"); st.stop()            
        # 1) Collect NEW PDFs
        pdf_files = list(Path(append_folder).glob("*.pdf"))
        if not pdf_files:
            st.warning("No PDFs found in the append folder!")
            st.stop()            
        # 2) Load existing artifacts
        try:
            index = faiss.read_index(exist_index_path)
            with open(exist_meta_path, "rb") as f:
                metadata_existing = pickle.load(f)
        except Exception as e:
            st.error(f"❌ Failed to load existing Knowledge-Base: {e}"); st.stop()
        # 2a) Verify embedding model compatibility
        existing_model = metadata_existing[0].get("embedding_model", st.session_state.embedding_model)
        st.session_state.embedding_model = existing_model
        st.session_state.dimension = d_em2dim[st.session_state.embedding_model]
        #
        # 3) Extract / clean / chunk new PDFs
        st.write("🔍 Processing NEW PDFs to append...")
        progress_bar = st.progress(0)
        new_chunks_by_file = {}
        all_texts = {}
        total_new_chunks = 0
        total_new_tokens = 0
        for i, pdf_path in enumerate(pdf_files):
            try:
                text = extract_text_from_pdf(pdf_path)
                text = remove_junk_sections(text)
                text = remove_junk_lines(text)
                all_texts[pdf_path.name] = text
                chunks = chunk_text2(text, max_tokens=TOKENS_PER_CHUNK, tokenizer=enc, overlap=WORDS_PER_CHUNK_OVERLAP)
                token_count = sum(len(enc.encode(c)) for c in chunks)
                new_chunks_by_file[pdf_path.name] = chunks
                total_new_chunks += len(chunks)
                total_new_tokens += token_count
            except Exception as e:
                st.warning(f"⚠️ Failed on {pdf_path.name}: {e}")
            progress_bar.progress(int((i+1)/len(pdf_files)*100))
        if total_new_chunks == 0:
            st.info("No new chunks were produced from the append PDFs."); st.stop()
        st.write(f"📦 New chunks: {total_new_chunks:,}, New tokens: {total_new_tokens:,}")
        #
        # 4) Embed NEW chunks and append to FAISS
        st.write("📌 Embedding and appending to FAISS...")
        new_embeddings = []
        new_metadata = []
        count = 0
        progress_bar = st.progress(0)
        for filename, chunks in new_chunks_by_file.items():
            for j, chunk in enumerate(chunks):
                try:
                    resp = client.embeddings.create(input=chunk, model=st.session_state.embedding_model)
                    emb = resp.data[0].embedding
                    new_embeddings.append(emb)
                    new_metadata.append({
                        "source": filename,
                        "chunk_id": j,
                        "text": chunk,
                        "embedding_model": st.session_state.embedding_model
                    })
                except Exception as e:
                    st.warning(f"⚠️ Embedding failed: {filename}, chunk {j} → {e}")
                count += 1
                progress_bar.progress(int(count / total_new_chunks * 100))
        if not new_embeddings:
            st.error("No new embeddings computed; aborting..."); st.stop()
        #
        new_mat = np.array(new_embeddings, dtype="float32")
        try:
            # Ensure dimension matches
            if new_mat.shape[1] != d_em2dim[st.session_state.embedding_model]:
                st.error("Embedding dimension mismatch. Aborting."); st.stop()
            index.add(new_mat)
        except Exception as e:
            st.error(f"Failed to append vectors to FAISS: {e}"); st.stop()

        # 5) Merge metadata and write outputs
        updated_metadata = metadata_existing + new_metadata
        os.makedirs(os.path.dirname(out_index_path), exist_ok=True)
        os.makedirs(os.path.dirname(out_meta_path), exist_ok=True)
        faiss.write_index(index, out_index_path)
        with open(out_meta_path, "wb") as f:
            pickle.dump(updated_metadata, f)
        #
        # 6) Rebuild LangChain FAISS wrapper + BM25 over all docs (existing + new)
        try:
            ids = [str(i) for i in range(len(updated_metadata))]
            docs_dict = {
                ids[i]: Document(
                    page_content=meta["text"],
                    metadata={
                        "source": meta["source"],
                        "chunk_id": meta["chunk_id"],
                        "original_content": meta.get("original_content", meta["text"]),
                    }
                )
                for i, meta in enumerate(updated_metadata)
            }
            docstore = InMemoryDocstore(docs_dict)
            index_to_docstore_id = {i: ids[i] for i in range(len(ids))}
            db = FAISS(
                embedding_function=embeddings_obj,
                index=index,
                docstore=docstore,
                index_to_docstore_id=index_to_docstore_id,
            )
            st.session_state.index = index
            st.session_state.metadata = updated_metadata            
            st.session_state.db = db
            all_documents = list(docstore._dict.values())
            st.session_state.all_documents = all_documents
            st.session_state.bm25 = BM25Retriever.from_documents(all_documents, k=top_k_textKBbm)
        except Exception as e:
            st.warning(f"KB wrapper rebuild warning: {e}")
            
        ## graphRAG
        if 'lightrag_engine' in st.session_state:
            kg_parent_dir = None
            kg_dir = None
            rag = None
            new_docs = None
            loop_append = None            
            del st.session_state.lightrag_engine               
        kg_parent_dir = Path(exist_knowledge_graph_parent)
        kg_dir = kg_parent_dir / "knowledge_graph" 
        required_files = [
            "graph_chunk_entity_relation.graphml", 
            "kv_store_doc_status.json", "kv_store_full_entities.json", "kv_store_full_relations.json", 
            "kv_store_full_docs.json", "kv_store_text_chunks.json", "kv_store_entity_chunks.json", "kv_store_relation_chunks.json", 
            "vdb_entities.json", "vdb_relationships.json", "vdb_chunks.json",
        ]
        missing = [f for f in required_files if not (kg_dir / f).exists()]
        if missing:
            st.error(f"❌ Cannot append. Knowledge-Graph is incomplete.\nMissing: {missing}")
            st.stop()
        rag = get_lightrag_engine(
            working_dir=str(kg_dir),
            api_key=st.session_state.api_key,
            embedding_model=st.session_state.embedding_model,
            embedding_dim=st.session_state.dimension,
        )
        
        new_docs = list(all_texts.values())
        # new_docs = [f"DOC_ID:{uuid.uuid4().hex}\n{text}" for text in all_texts.values()]        
        st.write("🕸️ Appending to existing Knowledge-Graph...")
        with st.spinner("🚀 Running pipeline..."):

            async def append_graph():
                await rag.initialize_storages()
                await rag.ainsert(new_docs)
                # await rag.finalize_storages()
                
            nest_asyncio.apply()
            loop_append = asyncio.new_event_loop()
            asyncio.set_event_loop(loop_append)
            try:
                loop_append.run_until_complete(append_graph())
            finally:
                # loop_append.close()     
                pass
                        
        st.session_state.lightrag_engine = rag            
        st.success("✅ Successfully appended new documents to the existing Knowledge-Graph!")
        ## graphRAG -
        st.success("✅ Append completed! Updated Index/Metadata/Graph have been saved!")            

st.divider()

# s3: Q&A
st.header("❓ Ask a Question")
st.session_state.gpt_model = st.selectbox(
    "🤖 Select model:",
    options=[
        "gpt-4.1-2025-04-14",                
        "gpt-4o-2024-08-06",        
        "gpt-5",
        "gpt-5-thinking",
        "gpt-5-pro",
        "o4-mini-2025-04-16",
        "o4-mini-deep-research-2025-06-26",
    ],
    index=0,
    help="This model generates the final answer!"
)

with st.expander("🎛️ Advanced Controls: Diversity & Creativity"):
    diversity = st.slider(
        "🧭 Diversity Radar (0 = Homogeneous, 1 = Diverse)",
        min_value=0.0,
        max_value=1.0,
        value=0.7,
        step=0.01,
        help="Higher values prioritize more varied sources; lower values focus more tightly!"
    )
    if st.session_state.gpt_model in ["gpt-4o-2024-08-06", "gpt-4.1-2025-04-14"]:
        temperature = st.slider(
            "🔥 Creativity Dial (0 = Boring, 1 = Creative)",
            min_value=0.0,
            max_value=1.0,
            value=0.3,
            step=0.01,
            help="Lower values are more deterministic!"
        )
    else:
        temperature = 0.3

if "index" in st.session_state:
    for key in ["last_query", "last_answer", "context_meta"]:
        if key not in st.session_state:
            st.session_state[key] = ""    
    query = st.text_area("Ask your question here:", height=280, placeholder="Type your question...")
    
    # Uploads UI kept from the extended app
    st.markdown("#### 📎 Add any relevant file(s) for this question (optional)")
    uploaded_files = st.file_uploader(
        "Upload PDFs/TXT/CSV/XLSX or images (PNG/JPG/JPEG/GIF/BMP/TIF/TIFF):",
        type=["pdf", "txt", "csv", "xlsx", "png", "jpg", "jpeg", "gif", "bmp", "tif", "tiff"],
        accept_multiple_files=True
    )
    use_uploads = True

    ## Answer
    # if st.button("💬 Answer") and query:
    col1, col2, col3 = st.columns([1, 2, 1])     
    with col1:
        answer_clicked = st.button("💬 Answer", use_container_width=True)
    with col3:
        save_clicked = st.button("💾 Save last Q&A", use_container_width=True)        
    st.markdown("")        
    if answer_clicked and query:
    ##  
        # 1) process uploads (text + images), cache for reuse
        with st.spinner("📂 Processing uploaded files..."):        
            if uploaded_files:
                try:
                    up_db, up_meta, up_images = build_upload_bundle(
                        uploaded_files=uploaded_files,
                        client=client,
                        embedding_model=st.session_state.embedding_model,
                        dimension=d_em2dim[st.session_state.embedding_model]
                    )
                    st.session_state.upload_db = up_db
                    st.session_state.upload_meta = up_meta
                    st.session_state.upload_images = up_images
                    n_text = len(up_meta)
                    n_imgs = len(up_images)
                    st.success(f"✅ Processed {n_text} text chunks and {n_imgs} image(s) from the uploaded files!")
                except Exception as e:
                    st.error(f"❌ Upload processing failed: {e}")
            else:
                st.session_state.upload_db = None
                st.session_state.upload_meta = []
                st.session_state.upload_images = []
    
            if use_uploads and "upload_meta" in st.session_state:
                st.caption(f"Uploads ready: {len(st.session_state.get('upload_meta', []))} text chunks, "
                           f"{len(st.session_state.get('upload_images', []))} images!")

        # 2) retrieve context (KB + optional uploads) via MMR
        with st.spinner("🔍 Knowledge-Base: Retrieving relevant context ..."):        
            query_embedding = client.embeddings.create(
                input=query,
                model=st.session_state.embedding_model
            ).data[0].embedding
            query_embedding = np.array(query_embedding, dtype="float32").reshape(1, -1)
            flat_emb = query_embedding[0].astype(float).tolist()
            ##
            ##
            # results = st.session_state.db.max_marginal_relevance_search_by_vector(
            #     embedding=flat_emb, k=top_k_textKBfaiss, fetch_k=top_k_textKBfaiss * 2, lambda_mult=1.0 - diversity
            # )
            vs_retriever = st.session_state.db.as_retriever(
                search_type="mmr",
                search_kwargs={
                    "k": top_k_textKBfaiss,
                    "fetch_k": top_k_textKBfaiss * 2,
                    "lambda_mult": 1.0 - diversity,
                },
            )
            bm25 = st.session_state.bm25
            hybrid = EnsembleRetriever(retrievers=[vs_retriever, bm25], weights=[0.5, 0.5])      
            llm_expander = ChatOpenAI(model="gpt-4o-mini")
            multi_query = MultiQueryRetriever.from_llm(
                retriever=hybrid,
                llm=llm_expander,
                include_original=True,
            ) 
            # cross_encoder = HuggingFaceCrossEncoder(model_name="cross-encoder/ms-marco-MiniLM-L-6-v2")
            # reranker = CrossEncoderReranker(model=cross_encoder, top_n=int(top_k_textKBfaiss+top_k_textKBbm))        
            # reranked_exapanded = ContextualCompressionRetriever(
            #     base_retriever=multi_query,
            #     base_compressor=reranker,
            # )        
            compressor_llm = ChatOpenAI(model="gpt-4o-mini", temperature=0)  
            compressor = LLMChainExtractor.from_llm(compressor_llm)    
            # compressed_retriever = ContextualCompressionRetriever(
            #     base_retriever=reranked_exapanded,
            #     base_compressor=compressor,
            # )        
            compressed_retriever = ContextualCompressionRetriever(
                base_retriever=multi_query,
                base_compressor=compressor,
            )                    
            results = compressed_retriever.invoke(query)
            reorder = LongContextReorder()
            results = reorder.transform_documents(results) 
            ##
            ## ==> till now (KB-text-faiss, KB-text-bm, KB-text-Engi)
            ## 
            results_up = []
            if use_uploads and st.session_state.get("upload_db") is not None:
                ##
                # results_up = st.session_state.upload_db.max_marginal_relevance_search_by_vector(
                #     embedding=flat_emb, k=top_k_addKBfaiss, fetch_k=top_k_addKBfaiss * 2, lambda_mult=1.0 - diversity
                # )
                #
                vs_retriever_up = st.session_state.upload_db.as_retriever(
                    search_type="mmr",
                    search_kwargs={
                        "k": top_k_addKBfaiss,
                        "fetch_k": top_k_addKBfaiss * 2,
                        "lambda_mult": 1.0 - diversity,
                    },
                )
                multi_query_up = MultiQueryRetriever.from_llm(
                    retriever=vs_retriever_up,
                    llm=llm_expander,
                    include_original=True,
                )
                compressed_retriever_up = ContextualCompressionRetriever(
                    base_retriever=multi_query_up,
                    base_compressor=compressor,
                )
                results_up = compressed_retriever_up.invoke(query)
                results_up = reorder.transform_documents(results_up)                
                ##
            merged = results        
            if results_up:
                merged.extend(results_up)
            context_meta_chunks = [
                f"[{doc.metadata['source']} | chunk {doc.metadata['chunk_id']}]: {doc.page_content}"
                for doc in merged
            ]        
            original_cmc = [
                f"[{doc.metadata['source']} | chunk {doc.metadata['chunk_id']}]: {doc.metadata['original_content']}"
                for doc in merged
            ]    
            ## ==> till now (upload-text-faiss, upload-text-Engi)
            ##
            if use_uploads and st.session_state.get("upload_images"):
                for img in st.session_state.upload_images[:]:
                    context_meta_chunks.append(f"[uploaded/{img['name']} | image]: (image attached)")  
                    original_cmc.append(f"[uploaded/{img['name']} | image]: (image attached)")  
            ## ==> till now (upload-image)
            ##
        ##
        with st.spinner("Knowledge-Graph: Retrieving entities, relationships, and summaries..."):        
            graph_context_chunks = []
            graph_context = ""
            original_gc = ""
            if "graphrag" in st.session_state and os.path.isdir(st.session_state.graphrag):
                try:
                    ROOT_DIR = Path(st.session_state.graphrag)
                    result = subprocess.run(
                        [
                            "graphrag", "query",
                            "--root", str(ROOT_DIR),
                            "--method", "global",
                            "--query", query
                        ],
                        capture_output=True,
                        text=True
                    )
                    if result.returncode !=0:
                        raise Exception(f"⚠️ Knowledge-Graph query failed:\n{result.stderr.strip()}")
                    answer = result.stdout.strip()
                    if answer:
                        # graph_context = (
                        #     "\n[GraphRAG Context]:\n" +
                        #     textwrap.shorten(answer, width=6000, placeholder=" ...")
                        # )
                        graph_context = (
                            "\n[Knowledge-Graph Context]:\n" + answer
                        )                        
                        context_meta_chunks.append(graph_context)
                        original_cmc.append(graph_context)
                        st.success("✅ Knowledge-Graph context successfully retrieved and added!")
                    else:
                        st.info(f"⚠️ No Knowledge-Graph response found for this query!")
                except Exception as e:
                    st.warning(f"⚠️ Error while retrieving from Knowledge-Graph: {e}")
            else:
                st.info("ℹ️ No Knowledge-Graph directory loaded - skipping graph-based retrieval.")
            ## ==> till now (KB-text-graph) 
        ##
        context_meta = "\n\n".join(context_meta_chunks)
        original_cm = "\n\n".join(original_cmc)
            
        # 3) build prompt & call model
        system_instructions = (
            "You are an expert scientific research assistant. Use the context provided from research papers to answer "
            "the user query as accurately as possible. Provide detailed responses using as much of the provided context "
            "as possible. If the answer is not clearly found in the context, respond with: "
            "'The context does not provide enough information to answer this question.' and default to your parametric knowledge to give a response "
            "If the provided context is not enough to answer the user, use web search to find relevant information. "
            "At the end of your answer, also mention the source file names or uploaded image names referenced in the context "
            "(e.g., [DL-rheed-harris-SI.pdf | chunk 1]) or (e.g., [uploaded/image.png | image])."
        )

        prompt_text = f"""
User query: {query}

--- BEGIN CONTEXT ---
{context_meta}
--- END CONTEXT ---

Answer:
""".strip()

        # Prepare image content blocks (if supported by model route)
        supports_images = st.session_state.gpt_model in ["gpt-4o-2024-08-06", "gpt-4.1-2025-04-14"]
        user_content = [{"type": "text", "text": prompt_text}]
        if use_uploads and st.session_state.get("upload_images") and supports_images:
            for img in st.session_state.upload_images[:]:
                user_content.append({"type": "text", "text": f"[uploaded/{img['name']} | image]:"})
                user_content.append({"type": "image_url", "image_url": {"url": img["data_url"]}})

        # Route per model
        model_choice = st.session_state.gpt_model

        try:
##
            st.markdown("### 💡 Answer")
            placeholder = st.empty()
            answer_text = ""
            
            if model_choice == "o4-mini-deep-research-2025-06-26":
                with client.responses.stream(
                    model=model_choice,
                    instructions=system_instructions,
                    tools=[{"type": "web_search_preview"}],
                    input=prompt_text,
                ) as stream:
                    for event in stream:
                        if event.type == "response.output_text.delta":
                            answer_text += event.delta
                            placeholder.markdown(answer_text + "▌")
                            time.sleep(stream_delay)  
                        elif event.type == "response.error":
                            st.error(str(event.error))
                        elif event.type in {"response.output_text.done", "response.completed"}:
                            placeholder.markdown(answer_text)     
                            
            elif model_choice in ["gpt-4o-2024-08-06", "gpt-4.1-2025-04-14"]:
                with client.chat.completions.stream(
                    model=model_choice,
                    messages=[
                        {"role": "system", "content": system_instructions},
                        {"role": "user", "content": user_content},
                    ],
                    temperature=temperature,
                ) as stream:
                    for event in stream:
                        if event.type == "content.delta":
                            answer_text += event.delta
                            placeholder.markdown(answer_text + "▌")
                            time.sleep(stream_delay)
                        elif event.type == "content.done":
                            placeholder.markdown(answer_text)
            
            elif model_choice in {"gpt-5", "gpt-5-thinking", "gpt-5-pro"}:
                model_id = "gpt-5" if model_choice == "gpt-5-thinking" else model_choice
                kwargs = {"model": model_id, "instructions": system_instructions, "input": prompt_text}
                if model_choice == "gpt-5-thinking":
                    kwargs["reasoning"] = {"effort": "high"}
                try:
                    with client.responses.stream(**kwargs) as stream:
                        for event in stream:
                            if event.type == "response.output_text.delta":
                                answer_text += event.delta
                                placeholder.markdown(answer_text + "▌")
                                time.sleep(stream_delay)
                            elif event.type == "response.error":
                                st.error(str(event.error))
                            elif event.type in {"response.output_text.done", "response.completed"}:
                                placeholder.markdown(answer_text)
                except Exception as e:
                    st.warning(f"⚠️ Streaming not supported for this model: {e}")
                    response = client.responses.create(**kwargs)
                    answer_text = response.output_text
                    placeholder.markdown(answer_text)            
            else:
                with client.chat.completions.stream(
                    model=model_choice,
                    messages=[
                        {"role": "system", "content": system_instructions},
                        {"role": "user", "content": prompt_text},
                    ],
                ) as stream:
                    for event in stream:
                        if event.type == "content.delta":
                            answer_text += event.delta
                            placeholder.markdown(answer_text + "▌")
                            time.sleep(stream_delay)
                        elif event.type == "content.done":
                            placeholder.markdown(answer_text)
## 
            # with st.expander("📚 Retrieved Context for this Query"):
            #     context_meta_html = original_cm.replace("\n", "<br>")
            #     # context_meta_html = context_meta.replace("\n", "<br>")                
            #     st.markdown(f"<div style='overflow-wrap: break-word; width: 600px'>{context_meta_html}</div>", unsafe_allow_html=True)

           ##
            st.session_state.last_query = query
            st.session_state.last_answer = answer_text
            st.session_state.context_meta = original_cm          
            # st.session_state.context_meta = context_meta                    
            ##            

        except Exception as e:
            st.error(f"❌ Inference failed: {e}")

    ##
    ## save
    if save_clicked and st.session_state.last_answer:
        try:
            save_path = Path("outputs/qa_pairs.docx")
            if save_path.exists():
                doc = DocxDocument(save_path)
            else:
                doc = DocxDocument()
                doc.add_heading("Question–Answer Log", level=1)
                doc.add_paragraph("")

            doc.add_heading("Question:", level=2)
            doc.add_paragraph(st.session_state.last_query.strip())
            doc.add_heading("Answer:", level=2)
            doc.add_paragraph(st.session_state.last_answer.strip())
            doc.add_paragraph("")  # spacing
            save_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(save_path)
            st.success(f"✅ Q&A pair added to: {save_path}")
        except Exception as e:
            st.error(f"❌ Saving Q&A pair failed: {e}")
    ##

    ## retrieved context
    if st.session_state.context_meta:
        st.divider()
        with st.expander("📚 Retrieved Context for this Query"):
            context_meta_html = st.session_state.context_meta.replace("\n", "<br>")            
            st.markdown(f"<div style='overflow-wrap: break-word; width: 600px'>{context_meta_html}</div>", unsafe_allow_html=True)
    ##   

else:
    st.info("⚠️ Please build or load a Knowledge-Base before asking a question!")
